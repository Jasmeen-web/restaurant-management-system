"""Generate the formatted DOCX report for the restaurant project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = PROJECT_ROOT / "report_assets" / "restaurant_report.md"
OUTPUT_PATH = PROJECT_ROOT / "Restaurant_Management_System_Report.docx"


def set_times_new_roman(run, size: int, bold: bool = False) -> None:
    """Apply Times New Roman formatting to a run."""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def justify_paragraph(paragraph) -> None:
    """Set paragraph alignment to justified."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_page_number(section) -> None:
    """Add a page number to the footer."""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_times_new_roman(run, 12)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def build_document() -> Document:
    """Build the report document from the markdown source."""
    document = Document()

    for section in document.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        add_page_number(section)

    lines = SOURCE_PATH.read_text(encoding="utf-8").splitlines()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[2:])
            set_times_new_roman(run, 16, bold=True)
            continue

        if stripped.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(stripped[3:])
            set_times_new_roman(run, 14, bold=True)
            continue

        if stripped.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(stripped[4:])
            set_times_new_roman(run, 13, bold=True)
            continue

        if stripped.startswith("[Insert Screenshot Here:"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped.strip("[]"))
            set_times_new_roman(run, 12, bold=True)
            run.italic = True
            continue

        paragraph = document.add_paragraph()
        justify_paragraph(paragraph)
        run = paragraph.add_run(stripped)
        set_times_new_roman(run, 12)

    return document


def count_words(text: str) -> int:
    """Return a simple word count."""
    return len([word for word in text.replace("\n", " ").split(" ") if word.strip()])


def main() -> None:
    """Generate the document and print the word count."""
    document = build_document()
    document.save(OUTPUT_PATH)
    source_text = SOURCE_PATH.read_text(encoding="utf-8")
    print(f"Report generated: {OUTPUT_PATH}")
    print(f"Approximate word count: {count_words(source_text)}")


if __name__ == "__main__":
    main()
